"""Extract text + a confidence signal from an uploaded creative (README §5).

Branches by what the file actually is (the extension does not tell you):
  * DOCX / text-PDF / txt   -> read the text layer directly
  * image / banner          -> vision pipeline (src.vision)
  * scanned / image PDF      -> rasterise pages, then vision
  * a folder of images       -> carousel: each frame through the vision pipeline

Two cheap defences against silently checking text the creative never contained:
extraction confidence is surfaced, and the extracted content is shown back.
"""
from __future__ import annotations

from pathlib import Path

from . import model, vision


def extract(path: str | Path) -> dict:
    p = Path(path)
    if p.is_dir():
        return _carousel(p)
    ext = p.suffix.lower()
    if ext == ".docx":
        return _docx(p)
    if ext == ".pdf":
        return _pdf(p)
    if ext in vision.IMAGE_EXTS:
        return _from_vision(vision.analyze_image(p), "image")
    if ext in {".txt", ".md"}:
        text = p.read_text(encoding="utf-8")
        return _result(text, "docx", 1.0 if text.strip() else 0.0)
    raise ValueError(f"Unsupported format {ext!r} (v1: .docx, .pdf, .txt, images, or a carousel folder)")


def _result(text: str, kind: str, conf: float, warnings=None) -> dict:
    return {
        "extracted_text": text,
        "source_kind": kind,
        "confidence": conf,
        "warnings": warnings or ([] if text.strip() else ["Extraction produced no text"]),
    }


def _from_vision(v: dict, kind: str, warnings=None) -> dict:
    """Wrap a vision result as an extraction dict; surface legibility as a warning."""
    warns = list(warnings or [])
    leg = v.get("disclaimer_legibility")
    if leg == "illegible":
        warns.append("Mandatory market-risk warning present but not legibly sized (vision)")
    elif leg == "absent":
        warns.append("Mandatory market-risk warning appears absent (vision)")
    return {
        "extracted_text": v.get("extracted_text", ""),
        "source_kind": kind,
        "confidence": float(v.get("confidence") or 0.0),
        "warnings": warns,
        "vision": v,
    }


def _docx(p: Path) -> dict:
    import docx  # python-docx

    d = docx.Document(str(p))
    parts = [para.text for para in d.paragraphs]
    for table in d.tables:
        for row in table.rows:
            parts.extend(cell.text for cell in row.cells)
    text = "\n".join(t for t in parts if t is not None)
    return _result(text, "docx", 0.99 if text.strip() else 0.0)


def _pdf(p: Path) -> dict:
    import pdfplumber

    with pdfplumber.open(str(p)) as pdf:
        pages = [page.extract_text() or "" for page in pdf.pages]
    text = "\n".join(pages).strip()
    if text:
        return _result(text, "pdf_text", 0.95)
    return _scanned_pdf(p)  # no text layer -> scanned / image-exported PDF


def _scanned_pdf(p: Path) -> dict:
    if not model.available():
        return _result("", "pdf_scanned", 0.0,
                       ["PDF has no text layer — scanned PDF needs the vision model (no API key set)"])
    import fitz  # pymupdf

    doc = fitz.open(str(p))
    frames = [vision.analyze_bytes(doc[i].get_pixmap(dpi=150).tobytes("png"), "image/png")
              for i in range(doc.page_count)]
    doc.close()
    return _from_vision(vision.merge_frames(frames), "pdf_scanned",
                        [f"scanned PDF — {len(frames)} page(s) read with vision"])


def _carousel(d: Path) -> dict:
    imgs = sorted(f for f in d.iterdir() if f.suffix.lower() in vision.IMAGE_EXTS)
    if not imgs:
        raise ValueError(f"No images found in carousel folder {d}")
    frames = [vision.analyze_image(f) for f in imgs]
    return _from_vision(vision.merge_frames(frames), "carousel", [f"carousel — {len(imgs)} frame(s)"])
